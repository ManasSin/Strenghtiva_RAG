"""
Per-collection vector store with OpenAI embeddings and cosine similarity retrieval.
Each document collection (dosha, diet, product, medical, lifestyle) gets its own index.
"""

import json
import os
from pathlib import Path
from typing import Any, Dict, List, Optional

import numpy as np
from openai import OpenAI


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _cosine(a: List[float], b: List[float]) -> float:
    va, vb = np.array(a, dtype=np.float32), np.array(b, dtype=np.float32)
    denom = np.linalg.norm(va) * np.linalg.norm(vb)
    return float(np.dot(va, vb) / (denom + 1e-10))


def _load_file_text(path: Path) -> str:
    suffix = path.suffix.lower()

    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")

    if suffix == ".pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            return "\n".join(p.extract_text() or "" for p in reader.pages)
        except ImportError:
            pass
        try:
            import pypdfium2 as pdfium
            doc = pdfium.PdfDocument(str(path))
            pages = []
            for page in doc:
                tp = page.get_textpage()
                pages.append(tp.get_text_range())
            return "\n".join(pages)
        except ImportError:
            return f"[PDF parsing unavailable — install pypdf: {path.name}]"

    if suffix == ".docx":
        try:
            from docx import Document
            doc = Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            return f"[DOCX parsing unavailable — install python-docx: {path.name}]"

    return ""


def _chunk(text: str, size: int = 500, overlap: int = 80) -> List[str]:
    words = text.split()
    chunks, i = [], 0
    while i < len(words):
        chunk = " ".join(words[i: i + size])
        if chunk.strip():
            chunks.append(chunk)
        i += size - overlap
    return chunks


# ---------------------------------------------------------------------------
# Per-collection store
# ---------------------------------------------------------------------------

class CollectionStore:
    INDEX_FILE = "_index.json"

    def __init__(self, name: str, folder: Path, client: OpenAI) -> None:
        self.name = name
        self.folder = folder
        self.client = client
        self._chunks: List[Dict[str, Any]] = []
        self._load()

    # ------------------------------------------------------------------
    def _index_path(self) -> Path:
        return self.folder / self.INDEX_FILE

    def _load(self) -> None:
        p = self._index_path()
        if p.exists():
            try:
                self._chunks = json.loads(p.read_text()).get("chunks", [])
            except Exception:
                self._chunks = []

    def _save(self) -> None:
        self._index_path().write_text(json.dumps({"chunks": self._chunks}, indent=2))

    # ------------------------------------------------------------------
    def ingest(self) -> int:
        """Re-index all documents in the folder. Returns chunk count."""
        if not self.folder.exists():
            return 0

        doc_files = [
            f for f in self.folder.iterdir()
            if f.suffix.lower() in {".pdf", ".docx", ".txt"}
            and f.name != self.INDEX_FILE
        ]
        if not doc_files:
            return 0

        raw_chunks: List[Dict[str, Any]] = []
        for doc in doc_files:
            text = _load_file_text(doc)
            if not text.strip():
                continue
            for idx, chunk_text in enumerate(_chunk(text)):
                raw_chunks.append({
                    "source": doc.name,
                    "chunk_id": f"{doc.stem}_{idx}",
                    "text": chunk_text,
                    "embedding": None,
                })

        # Batch embed (OpenAI allows up to 2048 inputs per request)
        BATCH = 100
        for i in range(0, len(raw_chunks), BATCH):
            batch = raw_chunks[i: i + BATCH]
            resp = self.client.embeddings.create(
                model="text-embedding-3-small",
                input=[c["text"] for c in batch],
            )
            for j, emb in enumerate(resp.data):
                batch[j]["embedding"] = emb.embedding

        self._chunks = raw_chunks
        self._save()
        return len(self._chunks)

    # ------------------------------------------------------------------
    def retrieve(self, query: str, top_k: int = 5) -> List[Dict]:
        """Return top_k chunks most similar to query."""
        if not self._chunks:
            return []

        q_emb = self.client.embeddings.create(
            model="text-embedding-3-small",
            input=[query],
        ).data[0].embedding

        scored = []
        for chunk in self._chunks:
            if chunk.get("embedding"):
                scored.append({**chunk, "score": _cosine(q_emb, chunk["embedding"])})

        scored.sort(key=lambda x: x["score"], reverse=True)
        return scored[:top_k]

    @property
    def chunk_count(self) -> int:
        return len(self._chunks)


# ---------------------------------------------------------------------------
# Multi-collection DB
# ---------------------------------------------------------------------------

class VectorDB:
    COLLECTIONS = ("dosha", "diet", "product", "medical", "lifestyle")

    def __init__(self, collection_paths: Dict[str, Path], api_key: str) -> None:
        self._api_key = api_key
        self._client: Optional[OpenAI] = None
        self.stores: Dict[str, CollectionStore] = {}
        for name in self.COLLECTIONS:
            path = collection_paths.get(name, Path(f"data/{name}_documents"))
            path.mkdir(parents=True, exist_ok=True)
            # Stores are created without a client; client is injected on first use
            self.stores[name] = CollectionStore.__new__(CollectionStore)
            self.stores[name].name = name
            self.stores[name].folder = path
            self.stores[name].client = None  # set lazily
            self.stores[name]._chunks = []
            self.stores[name]._load()

    @property
    def client(self) -> OpenAI:
        if self._client is None:
            key = self._api_key or os.getenv("OPENAI_API_KEY", "")
            if not key:
                raise ValueError("OPENAI_API_KEY is not set.")
            self._client = OpenAI(api_key=key)
        return self._client

    def _ensure_client(self, store: "CollectionStore") -> None:
        if store.client is None:
            store.client = self.client

    # ------------------------------------------------------------------
    def ingest_collection(self, name: str) -> int:
        if name not in self.stores:
            return 0
        self._ensure_client(self.stores[name])
        return self.stores[name].ingest()

    def get_chunk_count(self, name: str) -> int:
        return self.stores[name].chunk_count if name in self.stores else 0

    # ------------------------------------------------------------------
    def _get(self, name: str, query: str, top_k: int = 5) -> List[Dict]:
        if name not in self.stores:
            return []
        self._ensure_client(self.stores[name])
        return self.stores[name].retrieve(query, top_k=top_k)

    def retrieve_for_dosha(self, profile: Dict) -> List[Dict]:
        from features.dosha_analysis import build_dosha_query
        return self._get("dosha", build_dosha_query(profile), top_k=6)

    def retrieve_for_diet(self, profile: Dict) -> List[Dict]:
        from features.diet_recommendation import build_diet_query
        return self._get("diet", build_diet_query(profile), top_k=6)

    def retrieve_for_products(self, profile: Dict, top_k: int = 20) -> List[Dict]:
        from features.product_recommendation import build_product_query
        return self._get("product", build_product_query(profile), top_k=top_k)

    def retrieve_for_medical(self, profile: Dict) -> List[Dict]:
        diseases = [d for d in profile.get("existing_diseases", []) if d != "None"]
        q = f"Ayurvedic treatment diet recommendations for {', '.join(diseases) or 'general health'}"
        return self._get("medical", q, top_k=4)

    def retrieve_for_lifestyle(self, profile: Dict) -> List[Dict]:
        q = (
            f"Ayurvedic lifestyle daily routine recommendations for "
            f"{profile.get('physical_activity', 'moderate')} activity, "
            f"{profile.get('sleep_quality', 'average')} sleep quality, "
            f"{profile.get('stress_level', 'moderate')} stress level, "
            f"occupation: {profile.get('occupation', 'general')}"
        )
        return self._get("lifestyle", q, top_k=4)
